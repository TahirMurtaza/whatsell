"""
Knowledge Base Service — RAG (Retrieval Augmented Generation) pipeline.

Responsibilities:
- Document CRUD (backed by PostgreSQL)
- Text extraction from .txt / .pdf / .docx
- Simple chunking (500 chars, 50 overlap)
- Embedding via Gemini models/gemini-embedding-001 (768 dims)
- Cosine-similarity chunk retrieval via pgvector
- Strict-RAG chat: answers ONLY from uploaded documents
"""

import io
import logging
from typing import AsyncGenerator, List, Optional

import google.generativeai as genai
from sqlalchemy import select, text
from sqlalchemy.orm import Session as SyncSession

from app.config import get_settings
from app.db.postgres import async_session, sync_session
from app.models.postgres import Document, DocumentChunk

logger = logging.getLogger(__name__)
settings = get_settings()


# ---------------------------------------------------------------------------
# Helpers: text extraction
# ---------------------------------------------------------------------------

def _extract_text_plain(raw_bytes: bytes) -> str:
    return raw_bytes.decode("utf-8", errors="replace")


def _extract_text_from_pdf(raw_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw_bytes))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError(f"PDF extraction failed: {exc}") from exc


def _extract_text_from_docx(raw_bytes: bytes) -> str:
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(raw_bytes))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as exc:
        raise ValueError(f"DOCX extraction failed: {exc}") from exc


def _parse_document(raw_bytes: bytes, content_type: str) -> str:
    ct = (content_type or "").lower()
    if ct == "application/pdf" or ct.endswith("/pdf"):
        return _extract_text_from_pdf(raw_bytes)
    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_text_from_docx(raw_bytes)
    # Default: plain text (covers text/plain, text/*, unknown)
    return _extract_text_plain(raw_bytes)


# ---------------------------------------------------------------------------
# Helpers: chunking (simple sliding-window, no external dep)
# ---------------------------------------------------------------------------

def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks of ~chunk_size characters."""
    text = text.strip()
    if not text:
        return []
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


# ---------------------------------------------------------------------------
# Helpers: embedding (sync — used from Celery workers)
# ---------------------------------------------------------------------------

def _embed_texts_sync(texts: List[str]) -> List[List[float]]:
    """Embed texts one-at-a-time using EmbedContent RPC (not BatchEmbedContents).
    Passing a list to embed_content triggers BatchEmbedContents which has a hard
    60-second gRPC deadline and times out on large PDFs. Single-string calls each
    get their own 60-second budget and are immune to batch size issues."""
    import time
    genai.configure(api_key=settings.gemini_api_key)
    embeddings = []
    for i, text in enumerate(texts):
        result = genai.embed_content(
            model="models/gemini-embedding-001",
            content=text,
            task_type="retrieval_document",
        )
        embeddings.append(result["embedding"])
        # Brief pause every 10 calls to stay within Gemini rate limits
        if i > 0 and i % 10 == 0:
            time.sleep(0.5)
    return embeddings


def _embed_query_sync(query: str) -> List[float]:
    """Embed a query for similarity search (3072-dim vector)."""
    genai.configure(api_key=settings.gemini_api_key)
    result = genai.embed_content(
        model="models/gemini-embedding-001",
        content=query,
        task_type="retrieval_query",
    )
    return result["embedding"]


# ---------------------------------------------------------------------------
# Sync processing (called from Celery worker)
# ---------------------------------------------------------------------------

def process_document_sync(document_id: int, raw_bytes: bytes, content_type: str) -> None:
    """
    Parse → chunk → embed → store.  Runs inside a Celery worker (sync context).
    Updates document.status to 'ready' on success or 'error' on failure.
    """
    with sync_session() as db:
        doc = db.get(Document, document_id)
        if doc is None:
            logger.error(f"[KB] Document {document_id} not found")
            return

        doc.status = "processing"
        db.commit()

        try:
            # 1. Extract text
            full_text = _parse_document(raw_bytes, content_type)
            if not full_text.strip():
                raise ValueError("Document appears to be empty or unreadable.")

            # 2. Chunk
            chunks = _chunk_text(full_text)
            if not chunks:
                raise ValueError("No text chunks produced from document.")

            logger.info(f"[KB] Doc {document_id}: {len(chunks)} chunks, embedding now…")

            # 3. Embed in batches of 10 — each text is embedded individually
            # inside _embed_texts_sync to avoid gRPC BatchEmbedContents timeout
            BATCH = 10
            all_embeddings: List[List[float]] = []
            for i in range(0, len(chunks), BATCH):
                batch = chunks[i : i + BATCH]
                all_embeddings.extend(_embed_texts_sync(batch))

            # 4. Persist chunks
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, all_embeddings)):
                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=idx,
                    content=chunk_text,
                    embedding=embedding,
                )
                db.add(chunk)

            doc.status = "ready"
            db.commit()
            logger.info(f"[KB] Doc {document_id} ready ({len(chunks)} chunks)")

        except Exception as exc:
            logger.error(f"[KB] Doc {document_id} processing failed: {exc}", exc_info=True)
            doc.status = "error"
            doc.error_msg = str(exc)[:500]
            db.commit()
            raise  # re-raise so Celery can retry the task


# ---------------------------------------------------------------------------
# Async CRUD helpers (used from FastAPI endpoints)
# ---------------------------------------------------------------------------

async def create_document(
    session_id: Optional[str], filename: str, content_type: str
) -> Document:
    async with async_session() as db:
        doc = Document(
            session_id=session_id,
            filename=filename,
            content_type=content_type,
            status="pending",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        return doc


async def get_documents(session_id: str) -> List[dict]:
    async with async_session() as db:
        stmt = (
            select(Document)
            .where(Document.session_id == session_id)
            .order_by(Document.created_at.desc())
        )
        result = await db.execute(stmt)
        docs = result.scalars().all()
        return [
            {
                "id": d.id,
                "filename": d.filename,
                "content_type": d.content_type,
                "status": d.status,
                "error_msg": d.error_msg,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            }
            for d in docs
        ]


async def delete_document(document_id: int, session_id: str) -> bool:
    async with async_session() as db:
        stmt = select(Document).where(
            Document.id == document_id,
            Document.session_id == session_id,
        )
        result = await db.execute(stmt)
        doc = result.scalar_one_or_none()
        if doc is None:
            return False
        await db.delete(doc)
        await db.commit()
        return True


# ---------------------------------------------------------------------------
# RAG: chunk retrieval + chat
# ---------------------------------------------------------------------------

async def search_chunks(session_id: str, query: str, top_k: int = 5) -> List[dict]:
    """
    Embed the query then do cosine-similarity search against chunks owned by this session.
    Returns list of {content, filename, document_id}.
    """
    query_vec = _embed_query_sync(query)
    vec_literal = "[" + ",".join(str(v) for v in query_vec) + "]"

    sql = text(
        """
        SELECT dc.content, d.filename, dc.document_id,
               1 - (dc.embedding <=> CAST(:vec AS vector)) AS similarity
        FROM document_chunks dc
        JOIN documents d ON d.id = dc.document_id
        WHERE d.session_id = :session_id
          AND d.status = 'ready'
          AND dc.embedding IS NOT NULL
        ORDER BY dc.embedding <=> CAST(:vec AS vector)
        LIMIT :top_k
        """
    )

    async with async_session() as db:
        rows = await db.execute(
            sql,
            {"vec": vec_literal, "session_id": session_id, "top_k": top_k},
        )
        results = rows.fetchall()

    return [
        {"content": r[0], "filename": r[1], "document_id": r[2], "similarity": float(r[3])}
        for r in results
    ]


async def kb_chat(
    session_id: str,
    message: str,
    history: Optional[List[dict]] = None,
) -> AsyncGenerator[str, None]:
    """
    Strict RAG chat: retrieve top-5 relevant chunks, build context, stream Gemini reply.
    The LLM is explicitly forbidden from using outside knowledge.
    """
    # 1. Retrieve relevant chunks
    chunks = await search_chunks(session_id, message, top_k=5)

    if not chunks:
        yield "I don't have any documents in this session. Please upload some documents first, then ask your question."
        return

    # 2. Build context block
    context_parts = []
    for i, chunk in enumerate(chunks, 1):
        context_parts.append(
            f"[Source {i}: {chunk['filename']}]\n{chunk['content']}"
        )
    context_block = "\n\n---\n\n".join(context_parts)

    # 3. Build conversation history for Gemini
    history = history or []
    genai.configure(api_key=settings.gemini_api_key)
    model = genai.GenerativeModel(
        model_name=settings.gemini_model,
        system_instruction=(
            "You are a document assistant. You ONLY answer questions based on the "
            "provided document context below. You NEVER use outside knowledge or make "
            "up information. If the answer is not found in the context, say exactly: "
            "'I don't have that information in the uploaded documents.' "
            "Always cite which source document you used (e.g., 'According to [filename]...').\n\n"
            f"DOCUMENT CONTEXT:\n{context_block}"
        ),
    )

    # Build chat history
    gemini_history = []
    for msg in history:
        role = "user" if msg.get("role") == "user" else "model"
        gemini_history.append({"role": role, "parts": [msg.get("content", "")]})

    chat = model.start_chat(history=gemini_history)

    # 4. Stream response
    try:
        response = chat.send_message(message, stream=True)
        for chunk in response:
            if chunk.text:
                yield chunk.text
    except Exception as exc:
        logger.error(f"[KB] Gemini error: {exc}", exc_info=True)
        yield f"\n\nError generating response: {exc}"
